#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт генерации отчёта по ЛР №2 (локализация WPF) в формате ГОСТ.
Использование: python generate_report.py [ЛР_2_Фамилия_Имя.docx]
"""

import sys
import os

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Установите python-docx: pip install python-docx")
    sys.exit(1)


def set_gost_styles(doc):
    """Настройка стилей по ГОСТ (шрифт Times New Roman 14pt, межстрочный интервал 1.5)."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.first_line_indent = Cm(1.25)


def add_title_page(doc, author_name="Иванов Иван"):
    """Титульная страница по ГОСТ."""
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Министерство образования и науки Российской Федерации\n\n")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    run = title_para.add_run("Федеральное государственное бюджетное образовательное учреждение\nвысшего образования\n\n")
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    run = title_para.add_run("«Название университета»\n\n\n\n")
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    run = title_para.add_run("Лабораторная работа № 2\n")
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    run = title_para.add_run("по дисциплине «Объектно-ориентированное программирование»\n\n")
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    run = title_para.add_run("Тема: Локализация WPF-приложения\n")
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    run = title_para.add_run("Реализация многоязычной поддержки с использованием RESX, XAML ResourceDictionary и внешней библиотеки\n\n\n\n\n\n")
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"Выполнил: студент группы XXX\n{author_name}\n\nПроверил: _________________\n\n")
    r.font.size = Pt(14)
    r.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Город 2026")
    r.font.size = Pt(14)
    r.font.name = 'Times New Roman'

    doc.add_page_break()


def add_heading_gost(doc, text, level=1):
    """Заголовок с нумерацией по ГОСТ."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True
    return h


def add_paragraph_gost(doc, text):
    """Абзац с отступом первой строки."""
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
    return p


def main():
    output_file = sys.argv[1] if len(sys.argv) > 1 else "ЛР_2_Иванов_Иван.docx"
    author = "Иванов Иван"
    if len(sys.argv) > 2:
        author = sys.argv[2]

    doc = Document()
    set_gost_styles(doc)

    add_title_page(doc, author)

    add_heading_gost(doc, "1 Введение", level=1)
    add_paragraph_gost(doc,
        "Целью лабораторной работы является закрепление практических навыков локализации WPF-приложения. "
        "В работе изучаются и применяются три подхода к работе с переводами: использование файлов ресурсов (RESX), "
        "словарей ресурсов XAML (ResourceDictionary) и внешней библиотеки классов.")
    add_paragraph_gost(doc,
        "Задачи: модификация приложения из ЛР №1; добавление поддержки русского (по умолчанию) и английского языков; "
        "динамическое переключение языка без перезапуска приложения; реализация перевода статических элементов интерфейса "
        "и минимум одного строкового сообщения (например, в MessageBox) тремя указанными способами.")

    add_heading_gost(doc, "2 Теория", level=1)
    add_paragraph_gost(doc,
        "RESX (Resource XML) — файлы ресурсов .NET. Строки хранятся в XML, компилируются в сборку. "
        "Для разных культур создаются файлы вида Strings.resx (нейтральный/по умолчанию) и Strings.en.resx. "
        "ResourceManager.GetString(key, culture) возвращает строку для выбранной культуры.")
    add_paragraph_gost(doc,
        "Словари ресурсов XAML (ResourceDictionary) — ресурсы задаются в XAML (например, <sys:String x:Key=\"Key\">Значение</sys:String>). "
        "Для смены языка при динамическом переключении выполняется замена MergedDictionaries в ресурсах приложения.")
    add_paragraph_gost(doc,
        "Внешняя библиотека классов — отдельный проект (Class Library), содержащий строки по культурам (словари или RESX). "
        "Основное приложение ссылается на библиотеку и вызывает её API для получения строк по ключу и культуре.")

    add_heading_gost(doc, "3 Варианты перевода", level=1)

    add_heading_gost(doc, "3.1 Файлы ресурсов (RESX)", level=2)
    add_paragraph_gost(doc,
        "В ветке main реализована локализация через Properties/Strings.resx (русский по умолчанию) и Properties/Strings.en.resx (английский). "
        "Класс ResxLocalizationProvider использует ResourceManager с базовым именем View.Properties.Strings; при смене CurrentCulture вызывается OnPropertyChanged, "
        "и привязанные к свойствам провайдера элементы интерфейса обновляются. Сообщение в MessageBox при нажатии кнопки берётся через GetString(\"Message_DataSaved\").")
    add_paragraph_gost(doc, "[Место для скриншота: переключатель языка и интерфейс на русском и английском, MessageBox с переведённым текстом.]")

    add_heading_gost(doc, "3.2 Словари ресурсов XAML (ResourceDictionary)", level=2)
    add_paragraph_gost(doc,
        "В ветке xaml-dictionary строки вынесены в Localization/Strings.ru.xaml и Localization/Strings.en.xaml. "
        "XamlLocalizationProvider при смене культуры загружает нужный словарь по URI (pack://application:,,,/View;component/Localization/Strings.xx.xaml) "
        "и подменяет его в Application.Current.Resources.MergedDictionaries, затем уведомляет об обновлении. Интерфейс привязывается к тому же ключу Loc.")
    add_paragraph_gost(doc, "[Место для скриншота: работа переключения языка с XAML-словарями.]")

    add_heading_gost(doc, "3.3 Внешняя библиотека классов", level=2)
    add_paragraph_gost(doc,
        "В ветке external-library добавлен проект LocalizationLib (библиотека классов). Класс StringsProvider содержит статические словари Ru и En "
        "и метод GetString(CultureInfo, key). Приложение View ссылается на LocalizationLib; ExternalLibLocalizationProvider вызывает LocalizationLib.StringsProvider.GetString. "
        "Переключение языка и привязки работают аналогично остальным вариантам.")
    add_paragraph_gost(doc, "[Место для скриншота: интерфейс при использовании внешней библиотеки.]")

    add_heading_gost(doc, "4 Ссылка на проект в системе контроля версий", level=1)
    add_paragraph_gost(doc, "Репозиторий: https://github.com/<ваш-логин>/OOP_Lab_2")
    add_paragraph_gost(doc, "Ветки: main — RESX; xaml-dictionary — XAML ResourceDictionary; external-library — внешняя библиотека LocalizationLib.")

    add_heading_gost(doc, "5 Заключение", level=1)
    add_paragraph_gost(doc,
        "В ходе работы реализована многоязычная поддержка (русский и английский) в WPF-приложении тремя способами. "
        "Все варианты позволяют переключать язык без перезапуска. RESX удобен для интеграции с инструментами и сателлитными сборками; "
        "XAML ResourceDictionary хорошо сочетается с разметкой; внешняя библиотека даёт возможность переиспользовать строки в нескольких приложениях.")

    doc.save(output_file)
    print(f"Отчёт сохранён: {os.path.abspath(output_file)}")
    print("Замените «Название университета», группу, ФИО и добавьте скриншоты по меткам в документе.")


if __name__ == "__main__":
    main()
